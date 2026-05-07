"""
DOCX Exporter — Exporta entregables a Word (.docx).

Usa python-docx para generar documentos profesionales.
"""

import io
from typing import Any
import logging

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxExporter:
    """
    Exporta entregables a DOCX.

    Uso:
        exporter = DocxExporter(place_id, prospecto, params)
        filename, mimetype, bytes_data = exporter.export()
    """

    COLOR_VARKOS = RGBColor(0x16, 0x21, 0x3e)   # #16213e
    COLOR_ACCENT = RGBColor(0xe9, 0x45, 0x60)   # #e94560
    COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)    # #333

    def __init__(self, place_id: str, prospecto: dict, params: dict):
        self.place_id = place_id
        self.prospecto = prospecto
        self.params = params

    def export(self) -> tuple[str, str, io.BytesIO]:
        """
        Retorna (filename, mimetype, data_bytes).
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx no está instalado. Ejecuta: pip install python-docx")

        doc = Document()

        # Estilos base
        style = doc.styles["Normal"]
        style.font.name = "Georgia"
        style.font.size = Pt(11)
        style.font.color.rgb = self.COLOR_TEXT

        # Título
        nombre = self.prospecto.get("name", "Prospecto")
        p = doc.add_paragraph(nombre)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)
        p.runs[0].font.color.rgb = self.COLOR_VARKOS

        doc.add_paragraph("")

        # Contenido de la carta
        tipo = self.params.get("tipo_entregable", "carta_teaser")
        if tipo == "carta_teaser":
            self._add_carta_teaser(doc)
        else:
            self._add_texto_libre(doc)

        # Cierre Varkos
        doc.add_paragraph("")
        p_cierre = doc.add_paragraph(
            "Solo 20 espacios disponibles este mes para un análisis completo sin costo.\n"
            "También podemos ayudarles a equipar su clínica con las mejores marcas "
            "del mercado y financiamiento accesible."
        )
        p_cierre.runs[0].font.color.rgb = self.COLOR_TEXT

        doc.add_paragraph("")
        p_varkos = doc.add_paragraph("Varkos — Strategic Advice\nOscar\n[Teléfono/WhatsApp]")
        p_varkos.runs[0].font.color.rgb = self.COLOR_VARKOS
        p_varkos.runs[0].bold = True

        # Guardar a BytesIO
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe_name = nombre.replace(" ", "_").replace("/", "_")
        filename = f"Carta_Teaser_{safe_name}.docx"
        return filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", buf

    def _add_carta_teaser(self, doc):
        """Agrega contenido de carta teaser con bloques."""
        import re
        from cache_reader import CacheReader

        reader = CacheReader()
        editor_content = self.params.get("contenido_editor", "")

        # Cargar bloques: si hay marcadores en el editor, cargar todos
        editor_has_markers = bool(re.findall(r'\[bloque:\w+\]', editor_content))

        if editor_has_markers:
            all_bloques_meta = reader.list_bloques(self.place_id)
            bloques = {}
            for bmeta in all_bloques_meta:
                bid = bmeta.get("id")
                if bid:
                    full = reader.get_bloque(self.place_id, bid)
                    if full:
                        bloques[bid] = full

            def replace_bloque_marker(match):
                bloque_tipo = match.group(1)
                for bid, b in bloques.items():
                    if b.get("tipo") == bloque_tipo:
                        self._render_bloque_to_docx(doc, b)
                        return ""
                return ""

            # Procesar cada bloque del editor
            remaining = re.sub(
                r'\[bloque:(\w+)\](.*?)\[/bloque\]',
                replace_bloque_marker,
                editor_content,
                flags=re.DOTALL,
            )
            for para in remaining.split('\n\n'):
                p = para.strip()
                if p:
                    doc.add_paragraph(p)
        else:
            # Fallback: usar bloque_ids explícitos
            bloque_ids = self.params.get("bloque_ids", ["hook_dinero", "oportunidades"])
            for bid in bloque_ids:
                bloque = reader.get_bloque(self.place_id, bid)
                if bloque:
                    self._render_bloque_to_docx(doc, bloque)
                    doc.add_paragraph("")

    def _render_bloque_to_docx(self, doc, bloque: dict):
        """Renderiza un bloque como párrafos en el documento DOCX."""
        contenido = bloque.get("contenido") or {}
        tipo = bloque.get("tipo", "")

        if tipo == "hook_dinero":
            if isinstance(contenido, str):
                doc.add_paragraph(contenido)
            else:
                headline = contenido.get("headline", "")
                if headline:
                    p = doc.add_paragraph()
                    run = p.add_run(headline)
                    run.bold = True
                    run.font.size = Pt(13)
                    run.font.color.rgb = self.COLOR_ACCENT
                resumen = contenido.get("resumen", "")
                if resumen:
                    doc.add_paragraph(resumen)
                for d in contenido.get("datos_concretos", []):
                    if isinstance(d, dict):
                        label = d.get("label", "")
                        valor = d.get("valor", "")
                        doc.add_paragraph(f"• {label}: {valor}")

        elif tipo in ("oportunidades", "fortalezas", "comparativa_competitiva",
                      "insight_estrategico", "servicios_oportunidad"):
            if isinstance(contenido, str):
                text = contenido
            else:
                text = (
                    contenido.get("conclusion")
                    or contenido.get("ventaja")
                    or contenido.get("oportunidad")
                    or ""
                )
            if text:
                doc.add_paragraph(text)

        elif tipo == "temas_resenas":
            if isinstance(contenido, str):
                doc.add_paragraph(contenido)
            else:
                resumen = contenido.get("resumen", "")
                if resumen:
                    p = doc.add_paragraph()
                    run = p.add_run(resumen)
                    run.italic = True
                for t in contenido.get("temas_positivos", []):
                    doc.add_paragraph(f"✓ {t}")
                for t in contenido.get("temas_negativos", []):
                    doc.add_paragraph(f"✗ {t}")
                opp = contenido.get("oportunidad_principal", "")
                if opp:
                    p = doc.add_paragraph()
                    run = p.add_run(f"OPORTUNIDAD: {opp}")
                    run.bold = True

        else:
            if isinstance(contenido, str):
                text = contenido
            else:
                text = (
                    contenido.get("resumen")
                    or contenido.get("contenido")
                    or contenido.get("conclusion")
                    or ""
                )
            if text:
                doc.add_paragraph(text)

    def _add_texto_libre(self, doc):
        """Agrega texto libre del editor."""
        import re
        content = self.params.get("contenido_editor", "")
        if not content:
            return

        editor_has_markers = bool(re.findall(r'\[bloque:\w+\]', content))

        if editor_has_markers:
            # Ya manejado por _add_carta_teaser — aquí solo pasar texto residual
            return

        for para in content.split('\n\n'):
            p = para.strip()
            if p:
                doc.add_paragraph(p)